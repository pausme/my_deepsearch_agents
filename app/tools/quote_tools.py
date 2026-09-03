"""
装修报价单分析工具模块

封装报价单分析助手使用的四个 LangChain 工具：
parse_quote_file 解析报价单文件条目，normalize_quote_items 标准化归类，
compare_quote_with_reference_price 对比参考价区间，detect_quote_risk_items 做风险初筛。

文件解析只做确定性处理（表头识别、字段映射、金额提取），风险判断由
app/utils/quote_rules.py 的规则引擎和子智能体的模型判断共同完成。
"""

import json
import re
from pathlib import Path
from typing import Annotated

import pandas as pd
from langchain_core.tools import tool

from app.api.context import get_session_context
from app.api.monitor import monitor
from app.utils.path_utils import resolve_path
from app.utils.quote_rules import (
    compare_with_reference,
    detect_quote_risks,
    normalize_quote_items as normalize_items,
)

# 表头字段识别关键词：按优先级匹配列名
COLUMN_KEYWORDS = {
    "name": ["项目名称", "工程项目", "施工项目", "费用项目", "项目", "名称", "内容"],
    "unit": ["计量单位", "单位"],
    "quantity": ["工程量", "数量", "用量"],
    "unit_price": ["单价", "价格"],
    "total": ["合价", "金额", "小计", "总价"],
    "note": ["工艺说明", "材料说明", "品牌型号", "备注", "说明"],
}

# 合计行识别：行首单元格或行内出现这些字样时，把行内最大数字当作标注总价
TOTAL_ROW_KEYWORDS = ["总计", "合计", "报价总计", "总价合计", "报价合计", "优惠后总计"]

MAX_ROWS = 300  # 单个 sheet 最多解析条目数，防止异常大表拖垮上下文


def _match_column(header: str) -> str | None:
    """把报价单列名映射到标准字段；未命中返回 None。"""
    header_clean = re.sub(r"\s", "", str(header))
    if not header_clean:
        return None
    for field, keywords in COLUMN_KEYWORDS.items():
        for keyword in keywords:
            if keyword in header_clean:
                # “合价”包含“价”，优先匹配更长的关键词：合价/总价 先于 单价/价格
                if field == "unit_price" and any(k in header_clean for k in ["合价", "总价", "小计", "金额"]):
                    continue
                return field
    return None


def _find_header_row(rows: list[list]) -> int | None:
    """在前几行里找表头行：包含 >=2 个可识别字段的行。"""
    for index, row in enumerate(rows[:10]):
        matched = sum(1 for cell in row if cell is not None and _match_column(cell))
        if matched >= 2:
            return index
    return None


def _cell_str(value) -> str:
    """把单元格转成干净字符串；pandas 空单元格是 NaN，需要当成空串处理。"""
    if value is None:
        return ""
    if isinstance(value, float):
        if value != value:  # NaN 判定，避免引入 math 依赖
            return ""
        if value.is_integer():
            return str(int(value))
    return str(value).strip()


def _extract_items_from_rows(rows: list[list]) -> tuple[list[dict], str | None]:
    """
    从二维表格行中提取条目和标注总价

    :return: (条目列表, 标注总价字符串或 None)
    """
    header_index = _find_header_row(rows)
    items: list[dict] = []
    stated_total = None

    if header_index is None:
        # 无表头的自由文本：把非空行拼成原始文本交回给模型自行理解
        return items, stated_total

    header = rows[header_index]
    field_map: dict[int, str] = {}
    for col_index, cell in enumerate(header):
        field = _match_column(cell)
        if field and field not in field_map.values():
            field_map[col_index] = field

    for row in rows[header_index + 1 :]:
        cells = [_cell_str(cell) for cell in row]
        joined = "".join(cells)

        # 合计行：提取标注总价，不作为条目
        if any(keyword in joined for keyword in TOTAL_ROW_KEYWORDS):
            numbers = re.findall(r"[\d,]+(?:\.\d+)?", joined)
            if numbers:
                stated_total = max(numbers, key=lambda n: len(n)).replace(",", "")
            continue

        item: dict = {}
        for col_index, field in field_map.items():
            if col_index < len(cells):
                item[field] = cells[col_index]
        # 至少要有名称，且名称不能是纯数字，才认为是有效条目
        name = item.get("name", "")
        if name and not re.fullmatch(r"[\d.,\s]+", name):
            items.append(item)
        if len(items) >= MAX_ROWS:
            break

    return items, stated_total


def _parse_tabular_text(text: str) -> list[list]:
    """把 Markdown 管道表格 / CSV / 制表符文本解析成二维行，供条目提取复用。"""
    rows: list[list] = []
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        if line.startswith("|"):
            cells = [cell.strip() for cell in line.strip("|").split("|")]
            # 跳过 Markdown 表格分隔行 |---|---|
            if cells and all(re.fullmatch(r":?-{2,}:?", cell or "---") for cell in cells):
                continue
            rows.append(cells)
        elif "\t" in line or "," in line:
            splitter = "\t" if "\t" in line else ","
            rows.append([cell.strip() for cell in line.split(splitter)])
        else:
            rows.append([line])
    return rows


def _parse_quote_document(file_path: Path) -> dict:
    """按文件类型解析报价单，返回 {rows, raw_text, warnings}。"""
    ext = file_path.suffix.lower()
    warnings: list[str] = []

    if ext in [".xlsx", ".xls"]:
        sheets = pd.read_excel(str(file_path), sheet_name=None, header=None)
        rows: list[list] = []
        for sheet_name, df in sheets.items():
            sheet_rows = df.values.tolist()
            rows.append([f"# sheet: {sheet_name}"])
            rows.extend(sheet_rows)
        return {"rows": rows, "raw_text": "", "warnings": warnings}

    if ext in [".md", ".txt", ".csv"]:
        text = file_path.read_text(encoding="utf-8", errors="replace")
        return {"rows": _parse_tabular_text(text), "raw_text": text, "warnings": warnings}

    if ext == ".docx":
        try:
            import docx

            document = docx.Document(str(file_path))
            text = "\n".join(para.text for para in document.paragraphs if para.text.strip())
            # Word 中的报价表也尝试读取
            for table in document.tables:
                for row in table.rows:
                    text += "\n" + "|".join(cell.text.strip() for cell in row.cells)
            return {"rows": _parse_tabular_text(text), "raw_text": text, "warnings": warnings}
        except ImportError:
            warnings.append("未安装 python-docx，无法解析 Word 文件")
            return {"rows": [], "raw_text": "", "warnings": warnings}

    if ext == ".pdf":
        try:
            import pypdf

            reader = pypdf.PdfReader(str(file_path))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
            warnings.append("PDF 报价单按文本解析，扫描件或复杂版式可能无法提取表格")
            return {"rows": _parse_tabular_text(text), "raw_text": text, "warnings": warnings}
        except Exception as e:  # noqa: BLE001 - 工具层兜底，避免中断 Agent 链路
            warnings.append(f"PDF 解析失败：{e}")
            return {"rows": [], "raw_text": "", "warnings": warnings}

    warnings.append(f"暂不支持的报价单格式：{ext}，请转换为 Excel/PDF/Markdown 后重新上传")
    return {"rows": [], "raw_text": "", "warnings": warnings}


def _dump(data) -> str:
    """工具统一输出：紧凑 JSON 字符串（ensure_ascii=False 保证模型读到中文）。"""
    return json.dumps(data, ensure_ascii=False, default=str)


def _load_items(items_json: str) -> list[dict]:
    """宽容地解析上游工具传来的条目 JSON：接受数组或包装对象。

    兼容 parse_quote_file 的 {items: [...]} 和 normalize_quote_items 的
    {normalized_items: [...]}，模型直接把上一个工具的完整返回传进来也能工作。
    """
    data = json.loads(items_json)
    if isinstance(data, dict):
        data = data.get("items") or data.get("normalized_items") or []
    if not isinstance(data, list):
        raise ValueError("items_json 应为数组或包含 items/normalized_items 字段的对象")
    return data


@tool
def parse_quote_file(
    filename: Annotated[str, "报价单文件名（当前会话目录内，如 '装修报价单.xlsx'）"],
) -> str:
    """
    解析当前会话目录中的装修报价单文件，提取项目条目

    支持 Excel（多 sheet）、CSV、Markdown 表格、Word、PDF 文本。
    :param filename: 报价单文件名（不带目录前缀）
    :return: JSON 字符串，包含 items（条目：名称/单位/数量/单价/合价/备注）、
             stated_total（报价单标注总价）、warnings（解析提示）；解析失败时返回错误说明
    """
    monitor.report_tool("报价单解析工具", {"filename": filename})

    session_dir = get_session_context()
    file_path = Path(resolve_path(filename, session_dir))
    if not file_path.exists():
        return f"错误：文件 '{filename}' 不存在（解析路径: {file_path}），请先确认已上传。"

    parsed = _parse_quote_document(file_path)
    items, stated_total = _extract_items_from_rows(parsed["rows"])

    result = {
        "filename": filename,
        "item_count": len(items),
        "items": items,
        "stated_total": stated_total,
        "warnings": parsed["warnings"],
    }
    if not items and parsed["raw_text"]:
        # 表格识别失败时退回原文，让子智能体基于文本自行理解
        result["raw_text_head"] = parsed["raw_text"][:3000]
        result["warnings"].append("未能识别表格结构，已附原文前 3000 字，请基于原文分析条目")
    if not items and not parsed["raw_text"]:
        result["error"] = "未能从文件中解析出任何条目或文本"
    return _dump(result)


@tool
def normalize_quote_items(
    items_json: Annotated[str, "parse_quote_file 返回的 JSON（或其中的 items 数组）"],
) -> str:
    """
    标准化报价条目：清洗字段、按施工类目（拆改/水电/防水/瓦工/木作/油漆/安装/保洁清运）归类

    :param items_json: parse_quote_file 返回的 JSON 或 items 数组
    :return: JSON 字符串，包含 normalized_items（带 category 归类和 norm_key 去重键）
    """
    monitor.report_tool("报价条目标准化工具")

    items = _load_items(items_json)
    normalized = normalize_items(items)
    return _dump({"normalized_count": len(normalized), "normalized_items": normalized})


@tool
def compare_quote_with_reference_price(
    items_json: Annotated[str, "normalize_quote_items 返回的 JSON（或其中的 normalized_items 数组）"],
) -> str:
    """
    将报价条目单价与内置参考价区间对比，标记明显偏高或偏低的条目

    参考价为通用区间（二线城市口径），只作初筛提示，不构成实际报价依据。
    :param items_json: normalize_quote_items 返回的 JSON 或条目数组
    :return: JSON 字符串，包含 comparisons（条目价格 vs 参考区间）与 price_note（口径说明）
    """
    monitor.report_tool("报价参考价对比工具")

    items = _load_items(items_json)
    if items and "norm_key" not in items[0]:
        items = normalize_items(items)

    comparisons = compare_with_reference(items)
    flagged = [row for row in comparisons if row["flag"] != "ok"]
    return _dump(
        {
            "comparisons": comparisons,
            "flagged_count": len(flagged),
            "price_note": "参考价为通用区间，实际价格受城市、材料档次和工艺影响，仅供参考。",
        }
    )


@tool
def detect_quote_risk_items(
    items_json: Annotated[str, "normalize_quote_items 返回的 JSON（或其中的 normalized_items 数组）"],
    stated_total: Annotated[str, "报价单标注的合计金额（数字），未提取到时留空"] = "",
) -> str:
    """
    对标准化报价条目做风险初筛，输出结构化风险清单

    覆盖 7 类风险：疑似漏项、重复计费、计价单位含混、主材缺品牌型号、
    单价偏离参考区间、数量×单价与合价不一致、分项合计与总价不一致。
    :param items_json: normalize_quote_items 返回的 JSON 或条目数组
    :param stated_total: 报价单标注总价（可选，parse_quote_file 结果中的 stated_total）
    :return: JSON 字符串，包含 risks（每条含 risk_type/level/title/evidence/description/suggestion）
             和 summary（各级别风险计数）
    """
    monitor.report_tool("报价风险检测工具")

    items = _load_items(items_json)
    if items and "norm_key" not in items[0]:
        items = normalize_items(items)

    comparison = compare_with_reference(items)
    report = detect_quote_risks(items, comparison=comparison, stated_total=stated_total or None)

    # 高危风险单独上报前端，便于在执行轨迹中直接看到关键发现
    for risk in report["risks"]:
        if risk["level"] == "HIGH":
            monitor.report_risk(risk["level"], risk["title"], risk["description"])

    return _dump(report)
