"""
装修合同风险分析工具模块

封装合同风险助手使用的三个 LangChain 工具：
parse_contract_file 读取合同全文，extract_contract_clauses 提取关键条款，
match_contract_risk_rules 做规则化风险匹配。

条款提取和风险匹配逻辑在 app/utils/contract_rules.py，本模块负责文件读取、
monitor 埋点和工具输出格式。
"""

import json
from pathlib import Path
from typing import Annotated

from langchain_core.tools import tool

from app.api.context import get_session_context
from app.api.monitor import monitor
from app.utils.contract_rules import (
    extract_contract_clauses as extract_clauses_rule,
)
from app.utils.contract_rules import (
    match_contract_risk_rules as match_risk_rules_rule,
)
from app.utils.path_utils import resolve_path

MAX_CONTRACT_CHARS = 30000  # 合同全文超出时截断，防止上下文溢出


def _read_contract_text(file_path: Path) -> tuple[str, list[str]]:
    """按文件类型读取合同文本，返回 (文本, 警告列表)。"""
    ext = file_path.suffix.lower()
    warnings: list[str] = []

    if ext in [".md", ".txt"]:
        return file_path.read_text(encoding="utf-8", errors="replace"), warnings

    if ext == ".docx":
        try:
            import docx

            document = docx.Document(str(file_path))
            parts = [para.text for para in document.paragraphs if para.text.strip()]
            for table in document.tables:
                for row in table.rows:
                    row_text = " ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        parts.append(row_text)
            return "\n".join(parts), warnings
        except ImportError:
            warnings.append("未安装 python-docx，无法解析 Word 文件")
            return "", warnings

    if ext == ".pdf":
        try:
            import pypdf

            reader = pypdf.PdfReader(str(file_path))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
            if len(text.strip()) < 50:
                warnings.append("PDF 提取文本过少，可能是扫描件；建议转换为 Word 或文本后重新上传")
            return text, warnings
        except Exception as e:  # noqa: BLE001 - 工具层兜底，避免中断 Agent 链路
            warnings.append(f"PDF 解析失败：{e}")
            return "", warnings

    warnings.append(f"暂不支持的合同格式：{ext}，请转换为 PDF/Word/文本后重新上传")
    return "", warnings


def _dump(data) -> str:
    return json.dumps(data, ensure_ascii=False, default=str)


@tool
def parse_contract_file(
    filename: Annotated[str, "合同文件名（当前会话目录内，如 '装修合同.docx'）"],
) -> str:
    """
    读取当前会话目录中的装修合同文件，返回全文文本

    支持 docx、pdf、md、txt。
    :param filename: 合同文件名（不带目录前缀）
    :return: JSON 字符串，包含 text（合同全文，超长时截断）、char_count、warnings；
             解析失败时返回错误说明
    """
    monitor.report_tool("合同文件读取工具", {"filename": filename})

    session_dir = get_session_context()
    file_path = Path(resolve_path(filename, session_dir))
    if not file_path.exists():
        return f"错误：文件 '{filename}' 不存在（解析路径: {file_path}），请先确认已上传。"

    text, warnings = _read_contract_text(file_path)
    if not text and not warnings:
        return f"错误：文件 '{filename}' 内容为空。"

    char_count = len(text)
    truncated = char_count > MAX_CONTRACT_CHARS
    result = {
        "filename": filename,
        "char_count": char_count,
        "text": text[:MAX_CONTRACT_CHARS],
        "truncated": truncated,
        "warnings": warnings,
    }
    if truncated:
        result["warnings"] = warnings + [f"合同全文 {char_count} 字，已截断为前 {MAX_CONTRACT_CHARS} 字"]
    return _dump(result)


@tool
def extract_contract_clauses(
    contract_text: Annotated[str, "合同全文文本（parse_contract_file 返回的 text 字段）"],
) -> str:
    """
    从合同文本中提取关键条款要素

    提取内容：付款节点及各期比例、首期付款比例、竣工前累计付款、尾款比例、
    工期天数、延期违约、保修年限、增项确认流程、材料替换约定、违约与争议解决条款。
    :param contract_text: 合同全文文本
    :return: JSON 字符串，各条款要素；提取不到的字段为 null 或 false，请如实当作"未约定"
    """
    monitor.report_tool("合同条款提取工具")

    if not contract_text or not contract_text.strip():
        return _dump({"error": "合同文本为空，请先调用 parse_contract_file 读取合同"})

    clauses = extract_clauses_rule(contract_text)
    return _dump({"clauses": clauses, "notes": "提取不到的字段表示合同中未识别到对应约定"})


@tool
def match_contract_risk_rules(
    clauses_json: Annotated[str, "extract_contract_clauses 返回的 JSON"],
    contract_text: Annotated[str, "合同全文文本，用于给风险项附原文依据"] = "",
) -> str:
    """
    用内置规则对合同条款做风险初筛，输出结构化风险清单

    覆盖付款前置、竣工前累计过高、尾款过低、付款条款缺失、工期缺失或过短、
    延期违约缺失、增项无书面确认、材料替换无确认、保修缺失或不足、违约责任缺失等风险。
    :param clauses_json: extract_contract_clauses 返回的 JSON
    :param contract_text: 合同全文文本（可选，用于附原文依据）
    :return: JSON 字符串，包含 risks（每条含 risk_type/level/title/evidence/description/suggestion）、
             summary 和 disclaimer（免责声明）
    """
    monitor.report_tool("合同风险规则匹配工具")

    try:
        data = json.loads(clauses_json)
    except json.JSONDecodeError:
        return "错误：clauses_json 不是合法 JSON，请直接传入 extract_contract_clauses 的返回结果。"

    clauses = data.get("clauses", data) if isinstance(data, dict) else data
    if not isinstance(clauses, dict):
        return "错误：clauses_json 结构不符合预期。"

    report = match_risk_rules_rule(clauses, full_text=contract_text or "")

    for risk in report["risks"]:
        if risk["level"] == "HIGH":
            monitor.report_risk(risk["level"], risk["title"], risk["description"])

    # 风险项落库（FIX-006）：业务任务运行时关联到 renovation_risk_item
    try:
        from app.api.context import get_thread_context
        from app.repository import renovation_repository as repo

        thread_id = get_thread_context()
        if thread_id:
            repo.save_risk_items_by_thread(thread_id, report["risks"])
    except Exception:  # noqa: BLE001 - 持久化失败不影响 Agent 执行
        pass

    return _dump(report)
